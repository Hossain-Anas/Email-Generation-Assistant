"""Build the Final Report (Final_Report.docx) from the project artifacts.

The report contains the four required deliverable sections:
  1. The Prompt Template used (basic + advanced).
  2. The Definitions and Logic for the 3 Custom Metrics.
  3. The raw Evaluation Data (from results/evaluation_results.csv / .json).
  4. The Comparative Analysis summary (Section 3, from ANALYSIS.md).

Usage:
    python generate_report.py [--out Final_Report.docx]
"""
import argparse
import json

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from src.prompts import (
    BASIC_SYSTEM, BASIC_USER, ADVANCED_SYSTEM, ADVANCED_USER,
)

METRIC_DOC = [
    (
        "1. Fact Coverage  (LLM-judge, reference-free)",
        "Maps to the Key Facts input. The judge receives the email plus the numbered "
        "list of key facts and returns one boolean per fact. A fact counts as covered "
        "ONLY if its full meaning and every specific detail (names, numbers, dates, "
        "amounts, codes, IDs) is accurately present; a fact replaced by an unfilled "
        "placeholder or only vaguely alluded to is marked not covered.",
        "Score = (number of facts covered) / (total facts), in [0, 1]. Reference-free, so "
        "a well-reworded email is not penalized for differing from the human reference.",
    ),
    (
        "2. Tone Alignment  (LLM-judge)",
        "Maps to the Tone input. The judge rates how well the email matches the requested "
        "tone on a strict 1-5 rubric (5 = precisely and consistently matched; 3 = generic / "
        "partial; 1 = wrong tone). The rubric explicitly forbids defaulting to 5.",
        "Score = (rating - 1) / 4, normalizing 1->0.0 and 5->1.0.",
    ),
    (
        "3. Conciseness & Fluency  (hybrid: deterministic + LLM-judge)",
        "Maps to overall Quality and acts as a guardrail against the LLM-judge's known "
        "bias toward verbose answers. Combines a deterministic length component with an "
        "LLM grammar/fluency rating that penalizes filler, repetition, and unfilled "
        "placeholder fields.",
        f"length_score from word count (<= {config.LENGTH_FULL_MARKS_MAX} -> 1.0; "
        f"<= {config.LENGTH_PARTIAL_MAX} -> 0.7; else 0.4). "
        f"Final = {config.LENGTH_WEIGHT} * length_score + {config.FLUENCY_WEIGHT} * "
        "((fluency_rating - 1) / 4).",
    ),
]


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_markdown_ish(doc, md_text):
    """Render a subset of markdown (headings, bullets, tables, paragraphs) from a .md file."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # Markdown table block
        if line.lstrip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _render_md_table(doc, table_lines)
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(_strip_md(line.lstrip()[2:]), style="List Bullet")
        else:
            doc.add_paragraph(_strip_md(line))
        i += 1


def _strip_md(text):
    return text.replace("**", "").replace("`", "")


def _render_md_table(doc, table_lines):
    rows = [
        [c.strip() for c in ln.strip("|").split("|")]
        for ln in table_lines
        if "---" not in ln
    ]
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for c, val in enumerate(row):
            if c < len(cells):
                cells[c].text = _strip_md(val)
                if r == 0:
                    for p in cells[c].paragraphs:
                        for run in p.runs:
                            run.bold = True


def add_results_table(doc, df, columns, max_rows=None):
    cols = [c for c in columns if c in df.columns]
    sub = df[cols]
    if max_rows:
        sub = sub.head(max_rows)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for c, name in enumerate(cols):
        cell = table.rows[0].cells[c]
        cell.text = name
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for _, row in sub.iterrows():
        cells = table.add_row().cells
        for c, name in enumerate(cols):
            val = row[name]
            cells[c].text = f"{val:.3f}" if isinstance(val, float) else str(val)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="Final_Report.docx")
    args = parser.parse_args()

    csv_path = config.RESULTS_DIR / "evaluation_results.csv"
    json_path = config.RESULTS_DIR / "evaluation_results.json"
    analysis_path = config.ROOT / "ANALYSIS.md"

    df = pd.read_csv(csv_path)
    metric_cols = ["fact_coverage", "tone_alignment", "conciseness_fluency", "overall"]
    averages = df.groupby("strategy")[metric_cols].mean().round(3).reset_index()

    doc = Document()

    # ---- Title ----
    title = doc.add_heading("Email Generation Assistant", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Final Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    link = doc.add_paragraph("Repository: https://github.com/Hossain-Anas/Email-Generation-Assistant")
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generator: {config.GENERATOR_MODEL} ({config.GENERATOR_PROVIDER}) | "
        f"Judge: {config.JUDGE_MODEL} ({config.JUDGE_PROVIDER})"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== Section 1: Prompt Template =====
    add_heading(doc, "1. The Prompt Template", 1)
    doc.add_paragraph(
        "The assistant transforms Intent + Key Facts + Tone into a professional email. "
        "Two strategies are compared. The ADVANCED prompt is the graded technique; it "
        "layers three documented methods: Role-Playing (expert persona), silent "
        "Chain-of-Thought (internal planning that is not shown), and Few-Shot (one worked "
        "example), plus a strict output format lock."
    )

    add_heading(doc, "1a. Basic baseline", 2)
    doc.add_paragraph("System prompt:")
    add_code_block(doc, BASIC_SYSTEM)
    doc.add_paragraph("User prompt template:")
    add_code_block(doc, BASIC_USER)

    add_heading(doc, "1b. Advanced engineered prompt", 2)
    doc.add_paragraph("System prompt (Role-Playing):")
    add_code_block(doc, ADVANCED_SYSTEM)
    doc.add_paragraph("User prompt template (silent Chain-of-Thought + Few-Shot + format lock):")
    add_code_block(doc, ADVANCED_USER)

    # ===== Section 2: Metric Definitions & Logic =====
    add_heading(doc, "2. The 3 Custom Metrics: Definitions and Logic", 1)
    doc.add_paragraph(
        "All metrics are normalized to 0-1 and averaged into an overall score. The LLM "
        "judge runs at temperature 0 with JSON-only output; parsing is hardened against "
        "code fences, stray prose, and reasoning-model token exhaustion."
    )
    for name, definition, logic in METRIC_DOC:
        add_heading(doc, name, 2)
        p = doc.add_paragraph()
        p.add_run("Definition: ").bold = True
        p.add_run(definition)
        p = doc.add_paragraph()
        p.add_run("Logic: ").bold = True
        p.add_run(logic)

    # ===== Section 3: Raw Evaluation Data =====
    add_heading(doc, "3. Raw Evaluation Data", 1)
    add_heading(doc, "3a. Average scores by strategy", 2)
    add_results_table(doc, averages, ["strategy"] + metric_cols)

    add_heading(doc, "3b. Per-scenario scores (all 10 scenarios x 2 strategies)", 2)
    add_results_table(
        doc, df,
        ["id", "strategy", "fact_coverage", "tone_alignment", "conciseness_fluency", "overall"],
    )
    doc.add_paragraph(
        "Full machine-readable data, including each generated email and the judge's "
        "per-metric notes, is in results/evaluation_results.csv and "
        "results/evaluation_results.json."
    )

    # Sample generated emails (best advanced vs. its basic counterpart)
    try:
        with open(json_path, encoding="utf-8") as f:
            records = json.load(f).get("records", [])
        sample = next((r for r in records if r["id"] == "S08"), None)
        if sample is None and records:
            sample = records[0]
        if sample:
            sid = sample["id"]
            pair = {r["strategy"]: r for r in records if r["id"] == sid}
            add_heading(doc, f"3c. Sample output comparison (scenario {sid})", 2)
            for strat in ("basic", "advanced"):
                if strat in pair:
                    rec = pair[strat]
                    h = doc.add_paragraph()
                    h.add_run(
                        f"{strat.upper()} - overall {rec['overall']:.3f} "
                        f"(facts {rec['fact_coverage']:.2f}, tone {rec['tone_alignment']:.2f}, "
                        f"concise {rec['conciseness_fluency']:.2f}):"
                    ).bold = True
                    add_code_block(doc, rec["generated_email"])
    except (FileNotFoundError, json.JSONDecodeError, KeyError):
        pass

    # ===== Section 4: Comparative Analysis =====
    add_heading(doc, "4. Comparative Analysis Summary", 1)
    if analysis_path.exists():
        analysis_md = analysis_path.read_text(encoding="utf-8")
        # Drop the file's own top-level title to avoid a duplicate heading.
        analysis_md = "\n".join(
            ln for ln in analysis_md.splitlines()
            if not ln.startswith("# Comparative Analysis")
        )
        add_markdown_ish(doc, analysis_md)
    else:
        doc.add_paragraph("ANALYSIS.md not found.")

    out_path = config.ROOT / args.out
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
